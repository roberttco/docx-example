import docx
import pandas as pd

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX

from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

df = pd.read_excel("data_dictionary.xlsx")

# Open an existing Word document or create a new one
document = docx.Document("docx-example-tamplate.docx")

# Add a heading to the document
document.add_heading('DataFrame to Word Document', level=1)

# Add a table to the document
table = document.add_table(rows=df.shape[0] + 1, cols=df.shape[1])

styles = document.styles

# Set the header row (make it bold and center-aligned)
for j in range(df.shape[1]):
    cell = table.cell(0, j)
    cell.text = df.columns[j]

    # set the table cell background color
    # the color parse_xml call must be invoked for each cell
    hot_magenta = parse_xml(r'<w:shd {} w:fill="ff1dce"/>'.format(nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(hot_magenta)
    
    # get the paragraph that inside the table call
    paragraph = cell.paragraphs[0]

    # the table-heading style is applied to each heading cell paragraph
    paragraph.style = styles['table-heading']

    run = cell.paragraphs[0].runs[0]
    run.bold = True
    

# Populate the rest of the table with DataFrame values
for i in range(df.shape[0]):
    for j in range(df.shape[1]):
        cell = table.cell(i + 1, j)
        cell.text = str(df.values[i, j])

        if j == 0:
            # color the first column's cell backgrounds chartreuse
            chartreuse = parse_xml(r'<w:shd {} w:fill="7fff00"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(chartreuse)
        elif i % 2 == 0:
            # color every other row's cell backgrounds with sky blue
            light_sky_blue = parse_xml(r'<w:shd {} w:fill="87cefa"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(light_sky_blue)


        # get the paragraph that inside the table call
        paragraph = cell.paragraphs[0]

        # assign the styles to the cells.  The styles exist in the tempalte document
        # and control all aspects fo the text through the style definition.
        if j == 0:
            # the table-column1 style says right align
            paragraph.style = styles['table-column1']
        else:
            # if this were applied in addition to the table-colun1 style
            # the first column would be left aligned.
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Save the document
document.save('docx-example.docx')
